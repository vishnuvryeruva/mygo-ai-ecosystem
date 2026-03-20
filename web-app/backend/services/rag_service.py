import os
import io
import zipfile
import re

import PyPDF2
import psycopg2
import psycopg2.extras
from pgvector.psycopg2 import register_vector
from docx import Document

from db import get_conn
from services.openai_service import OpenAIService

# Supported file extensions for extraction from ZIP files
SUPPORTED_EXTENSIONS = {
    # Documents
    '.pdf', '.docx', '.txt', '.md',
    # Code files
    '.py', '.js', '.ts', '.java', '.abap', '.json', '.yaml', '.yml', '.xml'
}


class RAGService:
    def __init__(self):
        self.openai_service = OpenAIService()
        print("DEBUG: RAG Service initialized with OpenAI embeddings + PostgreSQL (pgvector)")

    # ── Embedding ─────────────────────────────────────────────────────────────

    def _create_embedding(self, text):
        """Create embedding using OpenAI's text-embedding-3-small model (1536-dim)."""
        try:
            from openai import OpenAI
            client = OpenAI(
                api_key=os.getenv('OPENAI_API_KEY'),
                timeout=120.0
            )
            response = client.embeddings.create(
                model="text-embedding-3-small",
                input=text
            )
            return response.data[0].embedding
        except Exception as e:
            print(f"ERROR: Failed to create embedding: {e}")
            raise

    # ── Internal helpers ───────────────────────────────────────────────────────

    def _chunk_text(self, text, chunk_size=500, overlap=50):
        """Split text into overlapping word-based chunks."""
        words = text.split()
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
        return chunks

    def _get_file_type(self, filename):
        """Map filename extension to a human-readable type label."""
        ext = filename.lower().split('.')[-1] if '.' in filename else 'unknown'
        type_map = {
            'pdf': 'PDF', 'docx': 'Word', 'doc': 'Word', 'txt': 'Text',
            'md': 'Markdown', 'py': 'Python', 'js': 'JavaScript',
            'ts': 'TypeScript', 'java': 'Java', 'abap': 'ABAP',
            'json': 'JSON', 'yaml': 'YAML', 'yml': 'YAML', 'xml': 'XML'
        }
        return type_map.get(ext, ext.upper())

    def _delete_chunks_for_doc(self, doc_id: str, conn):
        """Delete all existing chunks that start with doc_id."""
        with conn.cursor() as cur:
            cur.execute(
                "DELETE FROM documents WHERE id LIKE %s",
                (f"{doc_id}_%",)
            )

    def _insert_chunk(self, conn, chunk_id: str, doc_id: str, content: str,
                      embedding, metadata: dict):
        """Insert a single chunk row into the documents table."""
        with conn.cursor() as cur:
            cur.execute(
                """
                INSERT INTO documents (
                    id, document_id, document_name, content, embedding,
                    source, doc_type, project, updated_by, updated_on,
                    web_url, is_placeholder, html_content,
                    uuid, display_id, project_id, scope_id
                ) VALUES (
                    %s, %s, %s, %s, %s,
                    %s, %s, %s, %s, %s,
                    %s, %s, %s,
                    %s, %s, %s, %s
                )
                ON CONFLICT (id) DO UPDATE SET
                    content = EXCLUDED.content,
                    embedding = EXCLUDED.embedding,
                    html_content = EXCLUDED.html_content,
                    updated_on = EXCLUDED.updated_on
                """,
                (
                    chunk_id,
                    doc_id,
                    metadata.get('document_name', ''),
                    content,
                    embedding,
                    metadata.get('source', 'File Upload'),
                    metadata.get('type', 'Unknown'),
                    metadata.get('project', 'N/A'),
                    metadata.get('updatedBy', 'System'),
                    metadata.get('updatedOn', 'N/A'),
                    metadata.get('webUrl', ''),
                    metadata.get('is_placeholder', False),
                    metadata.get('html_content', ''),
                    metadata.get('uuid', ''),
                    metadata.get('displayId', ''),
                    metadata.get('projectId', ''),
                    metadata.get('scopeId', ''),
                )
            )

    # ── Public API ─────────────────────────────────────────────────────────────

    def check_document_exists(self, doc_id: str) -> bool:
        """Check if any chunk exists for the given document_id prefix."""
        conn = get_conn()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    "SELECT 1 FROM documents WHERE id LIKE %s LIMIT 1",
                    (f"{doc_id}_%",)
                )
                return cur.fetchone() is not None
        except Exception as e:
            print(f"Error checking document existence: {e}")
            return False
        finally:
            conn.close()

    def check_duplicate(self, filename: str) -> bool:
        """Check if a document with the same filename already exists."""
        return self.check_document_exists(filename)

    def ingest_calm_document(self, doc_metadata, html_content: str):
        """
        Ingest a CALM document with real HTML content into the vector database.
        Strips HTML tags for embedding/search, stores raw HTML in metadata for display.
        Replaces any existing placeholder or chunks for this document.
        """
        doc_id = doc_metadata.get('id', 'unknown')
        meta = doc_metadata.get('metadata', doc_metadata)
        filename = meta.get('name', doc_metadata.get('name', 'unknown'))
        doc_type = meta.get('type') or meta.get('documentType') or 'Document'
        source = meta.get('source', 'CALM')
        project = meta.get('project', 'N/A')
        updated_by = meta.get('updatedBy', 'System')
        updated_on = str(meta.get('updatedOn', meta.get('lastModified', 'N/A')))
        web_url = meta.get('webUrl', '')
        uuid_val = meta.get('uuid', '')
        display_id = meta.get('displayId', '')
        project_id = meta.get('projectId', '')
        scope_id = meta.get('scopeId', '')

        # Strip HTML tags to get plain text for embeddings
        plain_text = re.sub(r'<[^>]+>', ' ', html_content)
        plain_text = re.sub(r'\s+', ' ', plain_text).strip()

        if not plain_text:
            return self.add_placeholder_document(doc_metadata)

        MAX_HTML = 50000
        stored_html = html_content[:MAX_HTML] if len(html_content) > MAX_HTML else html_content

        base_metadata = {
            'source': source,
            'type': doc_type,
            'project': project,
            'updatedBy': updated_by,
            'updatedOn': updated_on,
            'webUrl': web_url,
            'is_placeholder': False,
            'document_name': filename,
            'uuid': uuid_val,
            'displayId': display_id,
            'projectId': project_id,
            'scopeId': scope_id,
        }

        chunks = self._chunk_text(plain_text, chunk_size=500, overlap=50)

        conn = get_conn()
        try:
            # Delete existing chunks for this document
            self._delete_chunks_for_doc(doc_id, conn)

            for i, chunk in enumerate(chunks):
                embedding = self._create_embedding(chunk)
                chunk_meta = dict(base_metadata)
                # Only store html_content on the first chunk
                chunk_meta['html_content'] = stored_html if i == 0 else ''
                self._insert_chunk(conn, f"{doc_id}_{i}", doc_id, chunk, embedding, chunk_meta)

            conn.commit()
        except Exception as e:
            conn.rollback()
            raise
        finally:
            conn.close()

        return {"status": "success", "chunks": len(chunks), "was_existing": False}

    def add_placeholder_document(self, doc_metadata):
        """Add or update a placeholder document (for synced external files with no content yet)."""
        try:
            meta = doc_metadata.get('metadata', doc_metadata)
            doc_id = doc_metadata.get('id', 'unknown')
            filename = meta.get('name', doc_metadata.get('name', 'unknown'))
            doc_type = meta.get('type') or meta.get('documentType') or 'Document'
            source = meta.get('source', 'CALM')
            project = meta.get('project', 'N/A')
            updated_by = meta.get('updatedBy', 'System')
            updated_on = meta.get('updatedOn', meta.get('lastModified', 'N/A'))
            web_url = meta.get('webUrl', '')
            uuid_val = meta.get('uuid', '')
            display_id = meta.get('displayId', '')
            project_id = meta.get('projectId', '')
            scope_id = meta.get('scopeId', '')

            already_exists = self.check_document_exists(doc_id)

            placeholder_text = (
                f"External document synced from {source}. Project: {project}. "
                f"Type: {doc_type}. Document: {filename}. "
                "This is a placeholder for metadata purposes."
            )
            embedding = self._create_embedding(placeholder_text)

            chunk_meta = {
                'source': source,
                'type': doc_type,
                'project': project,
                'updatedBy': updated_by,
                'updatedOn': str(updated_on),
                'webUrl': web_url,
                'is_placeholder': True,
                'document_name': filename,
                'uuid': uuid_val,
                'displayId': display_id,
                'projectId': project_id,
                'scopeId': scope_id,
                'html_content': '',
            }

            conn = get_conn()
            try:
                self._insert_chunk(conn, f"{doc_id}_0", doc_id, placeholder_text, embedding, chunk_meta)
                conn.commit()
            finally:
                conn.close()

            status = "updated" if already_exists else "success"
            return {"status": status, "message": "Document placeholder saved", "was_existing": already_exists}

        except Exception as e:
            print(f"Error adding/updating placeholder: {e}")
            return {"status": "error", "error": str(e)}

    def get_document_html(self, doc_id: str) -> str:
        """Retrieve the stored HTML content for a CALM document by its document_id.
        For File Upload docs (no html_content), returns concatenated chunk content wrapped in <pre>."""
        conn = get_conn()
        try:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    "SELECT html_content FROM documents WHERE document_id = %s AND html_content != '' AND html_content IS NOT NULL LIMIT 1",
                    (doc_id,)
                )
                row = cur.fetchone()
                if row and row['html_content']:
                    return row['html_content']
                # File Upload docs: no html_content, get chunk text
                cur.execute(
                    "SELECT content FROM documents WHERE document_id = %s ORDER BY id",
                    (doc_id,)
                )
                rows = cur.fetchall()
                if not rows:
                    return ''
                import html
                combined = "\n\n".join(r['content'] or '' for r in rows)
                return "<pre style='white-space:pre-wrap;font-family:inherit;'>" + html.escape(combined) + "</pre>"
        except Exception as e:
            print(f"Error retrieving html content for {doc_id}: {e}")
            return ''
        finally:
            conn.close()

    def ingest_documents(self, files, max_zip_size=256000):
        """Ingest uploaded files into the vector database."""
        results = []

        for file in files:
            try:
                filename_lower = file.filename.lower()

                # Handle ZIP files
                if filename_lower.endswith('.zip'):
                    file.seek(0, 2)
                    file_size = file.tell()
                    file.seek(0)

                    if file_size > max_zip_size:
                        results.append({
                            "filename": file.filename,
                            "status": "error",
                            "error": f"ZIP file exceeds 250KB limit ({file_size / 1024:.1f}KB)"
                        })
                        continue

                    zip_results = self._extract_zip(file)
                    results.extend(zip_results)
                    continue

                is_duplicate = self.check_duplicate(file.filename)
                text_content = self._extract_text(file)

                if not text_content.strip():
                    results.append({
                        "filename": file.filename,
                        "status": "error",
                        "error": "No text content extracted from file"
                    })
                    continue

                chunks = self._chunk_text(text_content, chunk_size=500, overlap=50)
                doc_id = file.filename

                conn = get_conn()
                try:
                    for i, chunk in enumerate(chunks):
                        embedding = self._create_embedding(chunk)
                        metadata = {
                            'source': 'File Upload',
                            'type': self._get_file_type(file.filename),
                            'project': 'N/A',
                            'updatedBy': 'System',
                            'updatedOn': 'N/A',
                            'document_name': file.filename,
                            'is_placeholder': False,
                            'html_content': '',
                        }
                        self._insert_chunk(conn, f"{doc_id}_{i}", doc_id, chunk, embedding, metadata)
                    conn.commit()
                except Exception:
                    conn.rollback()
                    raise
                finally:
                    conn.close()

                results.append({
                    "filename": file.filename,
                    "chunks": len(chunks),
                    "status": "success",
                    "was_duplicate": is_duplicate
                })

            except Exception as e:
                import traceback
                traceback.print_exc()
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": str(e)
                })

        return results

    def query(self, query_text, top_k=5, custom_prompt=None):
        """Query the RAG system using cosine similarity search.
        
        Returns:
            dict with 'answer' (str) and 'references' (list of source doc dicts)
        """
        query_embedding = self._create_embedding(query_text)

        conn = get_conn()
        try:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                cur.execute(
                    """
                    SELECT content, source, doc_type, project, document_name, web_url
                    FROM documents
                    ORDER BY embedding <=> %s::vector
                    LIMIT %s
                    """,
                    (query_embedding, top_k)
                )
                rows = cur.fetchall()
        finally:
            conn.close()

        # Build deduplicated references from matched chunks
        seen_names = set()
        references = []
        for row in rows:
            doc_name = row.get('document_name') or ''
            if doc_name and doc_name not in seen_names:
                seen_names.add(doc_name)
                references.append({
                    'document_name': doc_name,
                    'source': row.get('source', 'Unknown'),
                    'project': row.get('project', 'N/A'),
                    'doc_type': row.get('doc_type', 'Document'),
                    'web_url': row.get('web_url', ''),
                })

        context = "\n\n".join(row['content'] for row in rows)

        # Inject global document index
        try:
            doc_list = self.list_documents()
            global_index = "Available Documents Index:\n"
            for doc in doc_list:
                global_index += f"- '{doc['name']}' (Project: {doc['project']}, Source: {doc['source']})\n"
            context = global_index + "\n\nDetailed Excerpts:\n" + context
        except Exception as e:
            print(f"DEBUG: Failed to append global index to context: {e}")

        # Build system prompt
        if custom_prompt:
            system_prompt = custom_prompt
        else:
            try:
                from config.prompts import get_prompt
                system_prompt = get_prompt('ask_yoda', 'system')
                if not system_prompt:
                    system_prompt = (
                        "You are Yoda, a wise AI assistant with access to a knowledge base of SAP documents. "
                        "Answer questions based on the provided context. If the context doesn't contain enough "
                        "information, say so clearly."
                    )
            except Exception:
                system_prompt = (
                    "You are an expert SAP consultant and developer. "
                    "Answer questions based on the provided context from documents, test scripts, and tickets. "
                    "If the context doesn't contain enough information, say so clearly. "
                    "Provide clear, concise, and accurate answers."
                )

        user_prompt = (
            f"Context from knowledge base:\n{context}\n\n"
            f"Question: {query_text}\n\n"
            "Please provide a comprehensive answer based on the context above."
        )

        answer = self.openai_service.generate_text(
            user_prompt,
            system_prompt=system_prompt,
            temperature=0.3,
            max_tokens=1000
        )
        return {"answer": answer, "references": references}

    def list_documents(self):
        """List all unique documents with their metadata."""
        conn = get_conn()
        try:
            with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                # Get the first chunk per document_id (which holds html and full metadata)
                cur.execute(
                    """
                    SELECT DISTINCT ON (document_id)
                        document_id, document_name, source, doc_type, project,
                        updated_by, updated_on, web_url, uuid, display_id,
                        length(content) as content_len,
                        COUNT(*) OVER (PARTITION BY document_id) as chunk_count
                    FROM documents
                    ORDER BY document_id, id
                    """
                )
                rows = cur.fetchall()
        except Exception as e:
            print(f"Error listing documents: {e}")
            import traceback
            traceback.print_exc()
            return []
        finally:
            conn.close()

        doc_list = []
        for row in rows:
            estimated_size = (row['content_len'] or 0) * (row['chunk_count'] or 1)
            size_kb = estimated_size / 1024
            doc_list.append({
                'name': row['document_name'] or row['document_id'],
                'type': row['doc_type'] or 'Unknown',
                'source': row['source'] or 'File Upload',
                'project': row['project'] or 'N/A',
                'updatedBy': row['updated_by'] or 'System',
                'updatedOn': row['updated_on'] or 'N/A',
                'webUrl': row['web_url'],
                # Same value stored in chunks as document_id — required for view/delete/RAG keys
                'documentId': row['document_id'],
                'uuid': row['uuid'] or '',
                'size': f"{size_kb:.1f} KB" if size_kb >= 1 else f"{estimated_size} bytes",
                'chunks': row['chunk_count'] or 1,
            })

        return doc_list

    def delete_document(self, identifier: str) -> bool:
        """Delete all chunks for a document. Matches document_id (CALM id / upload filename) or legacy chunk id prefix."""
        if not identifier or not str(identifier).strip():
            return False
        ident = str(identifier).strip()
        conn = get_conn()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    DELETE FROM documents
                    WHERE document_id = %s
                       OR id LIKE %s
                       OR uuid::text = %s
                    """,
                    (ident, f"{ident}_%", ident),
                )
                deleted = cur.rowcount > 0
            conn.commit()
            return deleted
        except Exception as e:
            conn.rollback()
            print(f"Error deleting document: {e}")
            return False
        finally:
            conn.close()

    # ── ZIP extraction ─────────────────────────────────────────────────────────

    def _extract_zip(self, file):
        """Extract and process files from a ZIP archive."""
        results = []
        try:
            zip_data = io.BytesIO(file.read())
            with zipfile.ZipFile(zip_data, 'r') as zip_ref:
                for zip_info in zip_ref.filelist:
                    if zip_info.is_dir():
                        continue

                    inner_filename = zip_info.filename
                    _, ext = os.path.splitext(inner_filename.lower())

                    if ext not in SUPPORTED_EXTENSIONS:
                        continue

                    try:
                        file_content = zip_ref.read(inner_filename)
                        text_content = self._process_zip_content(inner_filename, file_content)

                        if not text_content.strip():
                            continue

                        base_name = os.path.basename(inner_filename)
                        is_duplicate = self.check_duplicate(base_name)
                        chunks = self._chunk_text(text_content, chunk_size=500, overlap=50)

                        conn = get_conn()
                        try:
                            for i, chunk in enumerate(chunks):
                                embedding = self._create_embedding(chunk)
                                metadata = {
                                    'source': 'File Upload',
                                    'type': self._get_file_type(base_name),
                                    'project': 'N/A',
                                    'updatedBy': 'System',
                                    'updatedOn': 'N/A',
                                    'document_name': base_name,
                                    'is_placeholder': False,
                                    'html_content': '',
                                }
                                self._insert_chunk(conn, f"{base_name}_{i}", base_name, chunk, embedding, metadata)
                            conn.commit()
                        except Exception:
                            conn.rollback()
                            raise
                        finally:
                            conn.close()

                        results.append({
                            "filename": f"{file.filename}/{inner_filename}",
                            "chunks": len(chunks),
                            "status": "success",
                            "was_duplicate": is_duplicate
                        })
                    except Exception as e:
                        results.append({
                            "filename": f"{file.filename}/{inner_filename}",
                            "status": "error",
                            "error": str(e)
                        })
        except zipfile.BadZipFile:
            results.append({"filename": file.filename, "status": "error", "error": "Invalid ZIP file"})
        except Exception as e:
            results.append({"filename": file.filename, "status": "error", "error": str(e)})

        return results

    def _process_zip_content(self, filename, content):
        """Process file content extracted from a ZIP archive."""
        filename_lower = filename.lower()
        if filename_lower.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
        elif filename_lower.endswith('.docx'):
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    return content.decode('latin-1')
                except Exception:
                    return ""

    def _extract_text(self, file):
        """Extract text from an uploaded file."""
        filename = file.filename.lower()
        if filename.endswith('.pdf'):
            return self._extract_pdf(file)
        elif filename.endswith('.docx'):
            return self._extract_docx(file)
        elif filename.endswith('.txt') or filename.endswith('.md'):
            return file.read().decode('utf-8')
        else:
            try:
                return file.read().decode('utf-8')
            except Exception:
                return ""

    def _extract_pdf(self, file):
        """Extract text from PDF."""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
        return "".join(page.extract_text() + "\n" for page in pdf_reader.pages)

    def _extract_docx(self, file):
        """Extract text from DOCX."""
        doc = Document(io.BytesIO(file.read()))
        return "\n".join(paragraph.text + "\n" for paragraph in doc.paragraphs)
