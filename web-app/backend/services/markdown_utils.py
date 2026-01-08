"""
Markdown to Word Document Utilities

This module provides utilities for converting markdown-formatted text
to properly styled Word document elements using python-docx.
"""

import re
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_rule(doc):
    """Add a horizontal rule (line) to the document"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    
    # Create a border element for the paragraph
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '999999')
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    return paragraph


def parse_inline_markdown(paragraph, text):
    """
    Parse inline markdown formatting and add runs to a paragraph.
    
    Handles:
    - **bold** or __bold__
    - *italic* or _italic_
    - `inline code`
    - Combined formatting
    """
    # Pattern to match markdown formatting
    # Order matters: check bold first (** or __), then italic (* or _), then code (`)
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)|'  # Bold with **
        r'(__(.+?)__)|'       # Bold with __
        r'(\*(.+?)\*)|'       # Italic with *
        r'(_([^_]+)_)|'       # Italic with _
        r'(`([^`]+)`)'        # Inline code
    )
    
    last_end = 0
    
    for match in pattern.finditer(text):
        # Add any text before this match as normal text
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])
        
        # Determine which group matched and apply formatting
        if match.group(2):  # Bold with **
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(4):  # Bold with __
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6):  # Italic with *
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8):  # Italic with _
            run = paragraph.add_run(match.group(8))
            run.italic = True
        elif match.group(10):  # Inline code
            run = paragraph.add_run(match.group(10))
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        
        last_end = match.end()
    
    # Add any remaining text after the last match
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def add_markdown_paragraph(doc, text):
    """
    Add a paragraph to the document, parsing markdown formatting.
    
    Handles:
    - Headings (# ## ### etc.)
    - Horizontal rules (--- or *** or ___)
    - Bullet lists (- or *)
    - Numbered lists (1. 2. etc.)
    - Inline formatting (bold, italic, code)
    
    Returns the created paragraph or None for special elements.
    """
    text = text.rstrip()
    
    if not text:
        return None
    
    # Check for horizontal rule
    if re.match(r'^(---+|\*\*\*+|___+)$', text.strip()):
        return add_horizontal_rule(doc)
    
    # Check for headings
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if heading_match:
        level = len(heading_match.group(1))
        heading_text = heading_match.group(2)
        # Clean any remaining markdown from heading
        heading_text = re.sub(r'\*\*(.+?)\*\*', r'\1', heading_text)
        heading_text = re.sub(r'__(.+?)__', r'\1', heading_text)
        return doc.add_heading(heading_text, level=min(level, 9))
    
    # Check for bullet list
    bullet_match = re.match(r'^[\-\*]\s+(.+)$', text)
    if bullet_match:
        paragraph = doc.add_paragraph(style='List Bullet')
        parse_inline_markdown(paragraph, bullet_match.group(1))
        return paragraph
    
    # Check for numbered list
    number_match = re.match(r'^(\d+)\.\s+(.+)$', text)
    if number_match:
        paragraph = doc.add_paragraph(style='List Number')
        parse_inline_markdown(paragraph, number_match.group(2))
        return paragraph
    
    # Regular paragraph with inline formatting
    paragraph = doc.add_paragraph()
    parse_inline_markdown(paragraph, text)
    return paragraph


def convert_markdown_to_docx(doc, content, title=None):
    """
    Convert markdown content to a Word document.
    
    Args:
        doc: A python-docx Document object
        content: The markdown content as a string
        title: Optional title to add as the document heading
    
    Returns:
        The modified document
    """
    if title:
        doc.add_heading(title, 0)
    
    # Split content into lines and process each
    lines = content.split('\n')
    
    # Track if we're in a code block
    in_code_block = False
    code_block_content = []
    
    for line in lines:
        # Check for code block markers
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block - add all collected content
                if code_block_content:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    for code_line in code_block_content:
                        run = code_para.add_run(code_line + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
        else:
            add_markdown_paragraph(doc, line)
    
    return doc
